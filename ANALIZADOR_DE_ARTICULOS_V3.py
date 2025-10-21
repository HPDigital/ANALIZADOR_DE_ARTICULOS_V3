"""
ANALIZADOR_DE_ARTICULOS_V3
"""

#!/usr/bin/env python
# coding: utf-8

# In[2]:


# from openai import OpenAI
# import fitz  # PyMuPDF

# client = OpenAI(api_key="YOUR_API_KEY_HERE")

# def extraer_texto_pdf(ruta_pdf):
#     # Abrir el documento PDF
#     documento = fitz.open(ruta_pdf)

#     # Recopilar el texto de cada página
#     texto_completo = ""
#     for pagina in documento:
#         texto_completo += pagina.get_text()

#     # Cerrar el documento
#     documento.close()
#     return texto_completo

# def analyze_article(texto_completo, client):
#     print("Iniciando el análisis del artículo científico.")

#     steps = [
#         ("titulo_articulo", "Cual es el titulo del artículo"),
#         ("Autor_articulo", "Quienes son los autores de este artículo"),
#         ("Objetivo_geenral_articulo", "Cual es el enunciado del objetivo general de este artículo"),
#         ("base teorica", "¿Cuáles son las bases teorícas sobre las cuales se basa el artículo de investigación?"),
#         ("Metodologia", "¿Cuál es la metodología con la que trabaja este artiíclo de investigación?"),
#         ("Herramientas", "¿Con que herraminetas de ivestigación se trabaja esta investigación"),
#         ("Muestra", "¿Que y cuantos utiliza como muestra para la investigacion"),
#         ("Resultados", "¿Que resultados se obtiene en la investigación?"),
#         ("Conclusiones", "Cuales son las conclusiones de esta ivestigación?"),
#         ("Conceptos_clave", " ¿Cuales son los conceptos clave que aborda el artículo?")
#     ]

#     for step_function, description in steps:
#         print(f"\n{description}")

#         response = client.chat.completions.create(
#             model="gpt-4o-2024-08-06", 
#             messages=[{"role": "system", "content": description},
#                       {"role": "user", "content": texto_completo}],
#             max_tokens=1024
#         )

#         # Asegurarse de que la respuesta tenga una estructura 
#         if hasattr(response, 'choices') and len(response.choices) > 0:
#             # Accede directamente al atributo 'content' del mensaje
#             message_content = response.choices[0].message.content.strip()
#             print(message_content)
#         else:
#             print("Error: No se obtuvieron resultados de la respuesta.")


# ruta_del_pdf = r"C:\Users\HP\Desktop\DIPLOMADO INVESTIGACION CIENTIFICA\MODULO 2\trabajo fin de modulo\ARTICULOS\Estudio de mercado del sector automotriz como herramienta para toma de decisiones empresariales.pdf"
# texto_completo = extraer_texto_pdf(ruta_del_pdf)
# analisis= analyze_article(texto_completo, client)
# print(analisis)



# In[ ]:





# In[3]:


# from openai import OpenAI
# import fitz  # PyMuPDF
# from docx import Document
# import os

# client = OpenAI(api_key="YOUR_API_KEY_HERE")

# def extraer_texto_pdf(ruta_pdf):
#     # Abrir el documento PDF
#     documento = fitz.open(ruta_pdf)

#     # Recopilar el texto de cada página
#     texto_completo = ""
#     for pagina in documento:
#         texto_completo += pagina.get_text()

#     # Cerrar el documento
#     documento.close()
#     return texto_completo

# def analyze_article(texto_completo, client):
#     print("Iniciando el análisis del artículo científico.")

#     steps = [
#         ("Título del artículo", "¿Cuál es el título del artículo?"),
#         ("Autores del artículo", "¿Quiénes son los autores de este artículo?"),
#         ("Objetivo general del artículo", "¿Cuál es el enunciado del objetivo general de este artículo?"),
#         ("Bases teóricas", "¿Cuáles son las bases teóricas sobre las cuales se basa el artículo de investigación?"),
#         ("Metodología", "¿Cuál es la metodología con la que trabaja este artículo de investigación?"),
#         ("Herramientas", "¿Con qué herramientas de investigación se trabaja en esta investigación?"),
#         ("Muestra", "¿Qué y cuántos utiliza como muestra para la investigación?"),
#         ("Resultados", "¿Qué resultados se obtienen en la investigación?"),
#         ("Conclusiones", "¿Cuáles son las conclusiones de esta investigación?"),
#         ("Conceptos clave", "¿Cuáles son los conceptos clave que aborda el artículo?")
#     ]

#     # Crear un nuevo documento Word
#     document = Document()

#     for step_name, description in steps:
#         print(f"\n{description}")

#         response = client.chat.completions.create(
#             model="gpt-4o-2024-08-06", 
#             messages=[{"role": "system", "content": description},
#                       {"role": "user", "content": texto_completo}],
#             max_tokens=1024
#         )

#         # Asegurarse de que la respuesta tenga una estructura 
#         if hasattr(response, 'choices') and len(response.choices) > 0:
#             # Accede directamente al atributo 'content' del mensaje
#             message_content = response.choices[0].message.content.strip()
#             print(message_content)

#             # Añadir el análisis al documento Word
#             document.add_heading(step_name, level=1)
#             document.add_paragraph(message_content)
#         else:
#             print("Error: No se obtuvieron resultados de la respuesta.")

#     return document

# def guardar_documento(document, ruta_pdf):
#     # Obtener el nombre del archivo PDF sin la extensión
#     nombre_archivo = os.path.splitext(os.path.basename(ruta_pdf))[0]

#     # Crear la ruta del archivo Word con el mismo nombre y en la misma carpeta
#     ruta_word = os.path.join(os.path.dirname(ruta_pdf), f"{nombre_archivo}.docx")

#     # Guardar el documento Word
#     document.save(ruta_word)
#     print(f"Documento guardado en: {ruta_word}")

# ruta_del_pdf = r"C:\Users\HP\Desktop\DIPLOMADO INVESTIGACION CIENTIFICA\MODULO 2\trabajo fin de modulo\ARTICULOS\Estudio de mercado del sector automotriz como herramienta para toma de decisiones empresariales.pdf"
# texto_completo = extraer_texto_pdf(ruta_del_pdf)
# document = analyze_article(texto_completo, client)
# guardar_documento(document, ruta_del_pdf)


# In[ ]:





# In[4]:


# ESTE FUNCIONA BIEN PARA ARTICULOS CORTOS

from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
import os

client = OpenAI(api_key="YOUR_API_KEY_HERE")

def extraer_texto_pdf(ruta_pdf):
    # Abrir el documento PDF
    documento = fitz.open(ruta_pdf)

    # Recopilar el texto de cada página
    texto_completo = ""
    for pagina in documento:
        texto_completo += pagina.get_text()

    # Cerrar el documento
    documento.close()
    return texto_completo

def analyze_article(texto_completo, client):
    print("Iniciando el análisis del artículo científico.")

    steps = [
        ("Título del artículo", "¿Cuál es el título del artículo?"),
        ("Autores del artículo", "¿Quiénes son los autores de este artículo?"),
        ("Objetivo general del artículo", "¿Cuál es el enunciado del objetivo general de este artículo?"),
        ("Bases teóricas", "¿Cuáles son las bases teóricas sobre las cuales se basa el artículo de investigación?"),
        ("Metodología", "¿Cuál es la metodología con la que trabaja este artículo de investigación?"),
        ("Herramientas", "¿Con qué herramientas de investigación se trabaja en esta investigación?"),
        ("Muestra", "¿Qué y cuántos utiliza como muestra para la investigación?"),
        ("Resultados", "¿Qué resultados se obtienen en la investigación?"),
        ("Conclusiones", "¿Cuáles son las conclusiones de esta investigación?"),
        ("Conceptos clave", "¿Cuáles son los conceptos clave que aborda el artículo?")
    ]

    # Crear un nuevo documento Word
    document = Document()

    for step_name, description in steps:
        #print(f"\n{description}")

        response = client.chat.completions.create(
            model="gpt-4o-2024-08-06", 
            messages=[{"role": "system", "content": description},
                      {"role": "user", "content": texto_completo}],
            max_tokens=1024
        )

        # Asegurarse de que la respuesta tenga una estructura 
        if hasattr(response, 'choices') and len(response.choices) > 0:
            # Accede directamente al atributo 'content' del mensaje
            message_content = response.choices[0].message.content.strip()
            #print(message_content)

            # Añadir el análisis al documento Word
            document.add_heading(step_name, level=1)
            document.add_paragraph(message_content)
        else:
            print("Error: No se obtuvieron resultados de la respuesta.")

    return document

def guardar_documento(document, ruta_pdf):
    # Obtener el nombre del archivo PDF sin la extensión
    nombre_archivo = os.path.splitext(os.path.basename(ruta_pdf))[0]

    # Crear la ruta del archivo Word con el mismo nombre y en la misma carpeta
    ruta_word = os.path.join(os.path.dirname(ruta_pdf), f"{nombre_archivo}.docx")

    # Guardar el documento Word
    document.save(ruta_word)
    print(f"Documento guardado en: {ruta_word}")

def procesar_todos_los_pdfs(directorio):
    # Listar todos los archivos PDF en el directorio
    archivos_pdf = [f for f in os.listdir(directorio) if f.endswith('.pdf')]

    for archivo_pdf in archivos_pdf:
        ruta_pdf = os.path.join(directorio, archivo_pdf)
        print(f"Procesando archivo: {ruta_pdf}")

        texto_completo = extraer_texto_pdf(ruta_pdf)
        document = analyze_article(texto_completo, client)
        guardar_documento(document, ruta_pdf)

# Ruta del directorio donde están los archivos PDF
directorio_pdfs = r"C:\Users\HP\Desktop\DIPLOMADO INVESTIGACION CIENTIFICA\MODULO 2\trabajo fin de modulo\ARTICULOS"
procesar_todos_los_pdfs(directorio_pdfs)


# In[6]:


from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
import os

client = OpenAI(api_key="YOUR_API_KEY_HERE")

def extraer_texto_pdf(ruta_pdf):
    documento = fitz.open(ruta_pdf)
    texto_completo = ""
    for pagina in documento:
        texto_completo += pagina.get_text()
    documento.close()
    return texto_completo

def analyze_article(texto_completo, client):
    print("Iniciando el análisis del artículo científico.")

    steps = [
        ("Título del artículo", "¿Cuál es el título del artículo?"),
        ("Autores del artículo", "¿Quiénes son los autores de este artículo?"),
        ("Objetivo general del artículo", "¿Cuál es el enunciado del objetivo general de este artículo?"),
        ("Bases teóricas", "¿Cuáles son las bases teóricas sobre las cuales se basa el artículo de investigación?"),
        ("Metodología", "¿Cuál es la metodología con la que trabaja este artículo de investigación?"),
        ("Herramientas", "¿Con qué herramientas de investigación se trabaja en esta investigación?"),
        ("Muestra", "¿Qué y cuántos utiliza como muestra para la investigación?"),
        ("Resultados", "¿Qué resultados se obtienen en la investigación?"),
        ("Conclusiones", "¿Cuáles son las conclusiones de esta investigación?"),
        ("Conceptos clave", "¿Cuáles son los conceptos clave que aborda el artículo?")
    ]

    document = Document()

    for step_name, pregunta in steps:
        try:
            response = client.chat.completions.create(
                model="gpt-4o-2024-08-06",
                messages=[{"role": "system", "content": pregunta},
                          {"role": "user", "content": texto_completo}],
                max_tokens=1024
            )

            if hasattr(response, 'choices') and len(response.choices) > 0:
                message_content = response.choices[0].message.content.strip()
                document.add_heading(step_name, level=1)
                document.add_paragraph(message_content)
            else:
                print("Error: No se obtuvieron resultados de la respuesta.")

        except Exception as e:
            print(f"Error durante el análisis de {step_name}: {e}")
            pass  # Pasar al siguiente paso o artículo en caso de error

    return document

def guardar_documento(document, ruta_pdf):
    nombre_archivo = os.path.splitext(os.path.basename(ruta_pdf))[0]
    ruta_word = os.path.join(os.path.dirname(ruta_pdf), f"{nombre_archivo}.docx")
    document.save(ruta_word)
    print(f"Documento guardado en: {ruta_word}")

def procesar_todos_los_pdfs(directorio):
    archivos_pdf = [f for f in os.listdir(directorio) if f.endswith('.pdf')]

    for archivo_pdf in archivos_pdf:
        ruta_pdf = os.path.join(directorio, archivo_pdf)
        print(f"Procesando archivo: {ruta_pdf}")

        try:
            texto_completo = extraer_texto_pdf(ruta_pdf)
            document = analyze_article(texto_completo, client)
            guardar_documento(document, ruta_pdf)
        except Exception as e:
            print(f"Error procesando el archivo {ruta_pdf}: {e}")
            pass  # Pasar al siguiente archivo en caso de error

# Ruta del directorio donde están los archivos PDF
directorio_pdfs = r"C:\Users\HP\Desktop\DIPLOMADO INVESTIGACION CIENTIFICA\MODULO 2\trabajo fin de modulo\ARTICULOS"
procesar_todos_los_pdfs(directorio_pdfs)


# In[ ]:






if __name__ == "__main__":
    pass
