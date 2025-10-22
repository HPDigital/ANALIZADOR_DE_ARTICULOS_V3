import json
import logging
import os
from typing import Callable, Dict, List

import fitz  # PyMuPDF
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI


load_dotenv()


SECTIONS: List[str] = [
    "Título del artículo",
    "Autores del artículo",
    "Objetivo general",
    "Bases teóricas",
    "Metodología",
    "Herramientas",
    "Muestra",
    "Resultados",
    "Conclusiones",
    "Conceptos clave",
]


def _ensure_logging() -> None:
    if not logging.getLogger().handlers:
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s %(levelname)s %(message)s",
            handlers=[logging.FileHandler("run.log", encoding="utf-8"), logging.StreamHandler()],
        )


def load_openai_client(api_key: str | None = None) -> OpenAI:
    key = api_key or os.getenv("OPENAI_API_KEY")
    if not key:
        raise ValueError(
            "Falta OPENAI_API_KEY. Define la variable de entorno o proporciónala en la interfaz."
        )
    return OpenAI(api_key=key)


def extract_text_from_pdf(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    try:
        parts = []
        for page in doc:
            parts.append(page.get_text())
        text = "".join(parts)
    finally:
        doc.close()
    return text.strip()


def _ocr_pdf_if_possible(pdf_path: str) -> str | None:
    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract  # type: ignore
    except Exception:
        return None
    try:
        images = convert_from_path(pdf_path)
        ocr_text = []
        for img in images:
            ocr_text.append(pytesseract.image_to_string(img, lang="spa+eng"))
        return "\n".join(ocr_text).strip()
    except Exception:
        return None


def _chunk_text(text: str, max_chars: int = 8000) -> List[str]:
    if len(text) <= max_chars:
        return [text]
    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        nl = text.rfind("\n", start, end)
        if nl != -1 and nl > start + 1000:
            end = nl
        chunks.append(text[start:end])
        start = end
    return chunks


def _approx_tokens(chars: int) -> int:
    return max(1, chars // 4)


def analyze_article(
    text: str,
    client: OpenAI,
    model: str = "gpt-4o-2024-08-06",
    temperature: float = 0.2,
    max_output_tokens: int = 1200,
    progress_cb: Callable[[str], None] | None = None,
    cancel_cb: Callable[[], bool] | None = None,
) -> Dict[str, str]:
    _ensure_logging()
    if cancel_cb and cancel_cb():
        raise RuntimeError("Operación cancelada por el usuario")

    chunks = _chunk_text(text, 8000)
    progress_cb and progress_cb(f"Texto dividido en {len(chunks)} fragmentos...")

    summaries: List[str] = []
    for i, ch in enumerate(chunks, 1):
        if cancel_cb and cancel_cb():
            raise RuntimeError("Operación cancelada por el usuario")
        if len(chunks) == 1:
            summaries = [ch]
            break
        prompt = (
            "Eres un asistente que resume contenido académico en español. "
            "Resume este fragmento preservando datos clave (métodos, muestras, resultados)."
        )
        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": ch},
        ]
        progress_cb and progress_cb(f"Pre-resumiendo fragmento {i}/{len(chunks)}...")
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=600,
            )
            summaries.append(resp.choices[0].message.content.strip())
        except Exception as e:
            logging.exception("Fallo en pre-resumen: %s", e)
            summaries.append(ch[:4000])

    merged_text = "\n\n".join(summaries)

    json_keys = {
        "Título del artículo": "string",
        "Autores del artículo": "string",
        "Objetivo general": "string",
        "Bases teóricas": "string",
        "Metodología": "string",
        "Herramientas": "string",
        "Muestra": "string",
        "Resultados": "string",
        "Conclusiones": "string",
        "Conceptos clave": "string",
    }

    sys_prompt = (
        "Eres un asistente experto en análisis de artículos científicos en español. "
        "Lee el contenido y devuelve UN ÚNICO JSON con claves exactamente como en el esquema. "
        "No incluyas texto fuera del JSON."
    )
    user_prompt = (
        "Analiza el contenido y completa las siguientes claves (breve y preciso):\n"
        f"Esquema: {list(json_keys.keys())}\n\n"
        f"Contenido:\n{merged_text}"
    )

    progress_cb and progress_cb("Solicitando análisis final (JSON)...")
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_prompt}],
            temperature=temperature,
            max_tokens=max_output_tokens,
        )
        raw = resp.choices[0].message.content.strip()
        start = raw.find("{")
        end = raw.rfind("}")
        if start != -1 and end != -1 and end > start:
            raw = raw[start : end + 1]
        data = json.loads(raw)
    except Exception as e:
        logging.exception("Fallo al parsear JSON: %s", e)
        data = {k: "" for k in SECTIONS}
        data["Conclusiones"] = f"No fue posible obtener JSON válido: {e}"

    results: Dict[str, str] = {k: str(data.get(k, "")).strip() for k in SECTIONS}
    return results


def build_docx(analysis: Dict[str, str], output_path: str) -> None:
    doc = Document()
    for section, content in analysis.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content or "Sin contenido")
    doc.save(output_path)


def process_pdf(
    pdf_path: str,
    client: OpenAI,
    out_dir: str | None = None,
    model: str = "gpt-4o-2024-08-06",
    temperature: float = 0.2,
    max_output_tokens: int = 1200,
    progress_cb: Callable[[str], None] | None = None,
    cancel_cb: Callable[[], bool] | None = None,
) -> str:
    _ensure_logging()
    progress_cb and progress_cb("Extrayendo texto del PDF...")
    text = extract_text_from_pdf(pdf_path)
    if len(text) < 50:
        progress_cb and progress_cb("Texto muy corto. Intentando OCR (si disponible)...")
        ocr_text = _ocr_pdf_if_possible(pdf_path)
        if ocr_text:
            text = ocr_text
        else:
            logging.info("OCR no disponible o falló. Continuando con texto detectado.")

    if cancel_cb and cancel_cb():
        raise RuntimeError("Operación cancelada por el usuario")

    analysis = analyze_article(
        text,
        client,
        model=model,
        temperature=temperature,
        max_output_tokens=max_output_tokens,
        progress_cb=progress_cb,
        cancel_cb=cancel_cb,
    )
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    target_dir = out_dir or os.path.dirname(pdf_path)
    os.makedirs(target_dir, exist_ok=True)
    out_path = os.path.join(target_dir, f"{base}.docx")
    progress_cb and progress_cb("Generando documento .docx...")
    build_docx(analysis, out_path)
    return out_path


def find_pdfs(directory: str, recursive: bool = True) -> List[str]:
    paths: List[str] = []
    if recursive:
        for root, _, files in os.walk(directory):
            for f in files:
                if f.lower().endswith(".pdf"):
                    paths.append(os.path.join(root, f))
    else:
        paths = [os.path.join(directory, f) for f in os.listdir(directory) if f.lower().endswith(".pdf")]
    return sorted(paths)

